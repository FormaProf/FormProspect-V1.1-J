from __future__ import annotations

import re
import shutil
import sqlite3
from core.sqlite_utils import connect_database
import subprocess
import tempfile
import zipfile
from lxml import etree
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover - handled at runtime with a clear message
    Document = None

from services.document_service import DocumentService


@dataclass(frozen=True)
class GeneratedDocument:
    document_id: int
    document_type: str
    document_number: str
    docx_path: Path
    pdf_path: Path | None
    template_name: str = "Modèle interne Form@Prof"
    warning: str = ""


@dataclass(frozen=True)
class GenerationPreflight:
    session_id: int
    errors: tuple[str, ...]
    warnings: tuple[str, ...]
    templates: dict[str, str]
    pdf_available: bool

    @property
    def can_generate(self) -> bool:
        return not self.errors


class DocumentGeneratorService:
    """RC2.0.3 — generation DOCX/PDF from a client training session.

    A configured DOCX template may contain variables such as {{CLIENT_NOM}}.
    When no suitable template is configured, a clean Form@Prof document is
    generated from scratch so generation never depends on a hard-coded client.
    """

    SUPPORTED_TYPES = ("Devis", "Convention", "Programme", "Convocation")
    UI_TYPES = ("Convention", "Programme", "Convocation")
    FOLDERS = {
        "Devis": "01 - Devis",
        "Convention": "02 - Conventions",
        "Programme": "03 - Programmes",
        "Convocation": "04 - Convocations",
    }

    def __init__(self, database_path: str | Path, project_folder: str | Path):
        self.database_path = Path(database_path)
        self.project_folder = Path(project_folder)
        self.document_service = DocumentService(self.database_path)
        self.document_service.ensure_schema()
        self.document_service.ensure_official_templates()

    def _connect(self) -> sqlite3.Connection:
        conn = connect_database(self.database_path)
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA foreign_keys = ON")
        return conn


    @staticmethod
    def pdf_converter_available() -> bool:
        """Retourne True si une conversion PDF locale est vraisemblablement possible."""
        return bool(shutil.which("libreoffice") or shutil.which("soffice") or shutil.which("powershell"))

    def preflight(self, session_id: int, document_types: Iterable[str] | None = None) -> GenerationPreflight:
        """Valide le dossier avant génération et décrit les modèles utilisés."""
        session = self.get_session_context(session_id)
        requested = list(dict.fromkeys(document_types or self.SUPPORTED_TYPES))
        errors: list[str] = []
        warnings: list[str] = []

        invalid = [item for item in requested if item not in self.SUPPORTED_TYPES]
        if invalid:
            errors.append("Type de document non pris en charge : " + ", ".join(invalid))
        if not requested:
            errors.append("Sélectionnez au moins un document à générer.")

        required = {
            "company_name": "le nom de l’entreprise",
            "training_name": "l’intitulé de la formation",
            "start_date": "la date de début",
            "end_date": "la date de fin",
            "modality": "la modalité",
        }
        for field, label in required.items():
            if not str(session.get(field) or "").strip():
                errors.append(f"Renseignez {label}.")
        try:
            start = date.fromisoformat(str(session.get("start_date") or ""))
            end = date.fromisoformat(str(session.get("end_date") or ""))
            if end < start:
                errors.append("La date de fin ne peut pas être antérieure à la date de début.")
        except ValueError:
            errors.append("Les dates de formation sont invalides.")
        if float(session.get("duration_hours") or 0) <= 0:
            errors.append("La durée de formation doit être supérieure à 0 heure.")
        if int(session.get("participant_count") or 0) <= 0:
            errors.append("Le nombre de participants doit être supérieur à 0.")

        optional = {
            "siret": "SIRET client non renseigné",
            "address": "adresse client non renseignée",
            "contact_name": "contact client non renseigné",
            "email": "adresse e-mail client non renseignée",
            "trainer_name": "formateur non renseigné",
            "daily_schedule": "horaires non renseignés",
        }
        for field, message in optional.items():
            if not str(session.get(field) or "").strip():
                warnings.append(message + ".")
        if int(session.get("price_cents") or 0) <= 0:
            warnings.append("Le tarif est nul : vérifiez la formation avant de générer le devis ou la convention.")

        templates: dict[str, str] = {}
        for document_type in requested:
            if document_type not in self.SUPPORTED_TYPES:
                continue
            template = self._find_template(document_type, int(session["training_id"]))
            if not template:
                if document_type == "Devis":
                    templates[document_type] = "Modèle interne Form@Prof"
                    warnings.append("Devis : aucun modèle DOCX associé, le modèle interne sera utilisé. Utilisez de préférence le workflow Demande de devis.")
                else:
                    templates[document_type] = "Aucun modèle"
                    errors.append(f"{document_type} : aucun modèle Word officiel actif n’est configuré.")
                continue
            source = Path(str(template.get("source_path") or ""))
            if source.suffix.lower() != ".docx":
                if document_type == "Devis":
                    templates[document_type] = "Modèle interne Form@Prof"
                    warnings.append("Devis : aucun modèle DOCX valide, le modèle interne sera utilisé. Le workflow Demande de devis est recommandé.")
                else:
                    templates[document_type] = "Modèle invalide"
                    errors.append(f"{document_type} : le modèle « {template.get('name', '')} » doit être un fichier DOCX.")
            elif not source.is_file():
                if document_type == "Devis":
                    templates[document_type] = "Modèle interne Form@Prof"
                    warnings.append("Devis : aucun modèle DOCX disponible, le modèle interne sera utilisé. Le workflow Demande de devis est recommandé.")
                else:
                    templates[document_type] = "Fichier introuvable"
                    errors.append(f"{document_type} : fichier modèle introuvable ({source}).")
            else:
                templates[document_type] = str(template.get("name") or source.name)

        return GenerationPreflight(
            session_id=int(session_id),
            errors=tuple(dict.fromkeys(errors)),
            warnings=tuple(dict.fromkeys(warnings)),
            templates=templates,
            pdf_available=self.pdf_converter_available(),
        )

    def list_sessions(self) -> list[dict[str, Any]]:
        with self._connect() as conn:
            rows = conn.execute(
                """SELECT s.id AS session_id, s.*, c.id AS client_id, c.prospect_id,
                          c.company_name, c.siret, c.address, c.postal_code, c.city,
                          c.contact_name, c.phone, c.email, c.folder_path,
                          t.reference AS training_reference, t.name AS training_name,
                          t.description AS training_description, t.objectives,
                          t.prerequisites, t.certification
                   FROM training_sessions s
                   JOIN clients c ON c.id=s.client_id
                   JOIN trainings t ON t.id=s.training_id
                   ORDER BY s.start_date DESC, c.company_name COLLATE NOCASE"""
            ).fetchall()
        return [dict(row) for row in rows]

    def get_session_context(self, session_id: int) -> dict[str, Any]:
        with self._connect() as conn:
            row = conn.execute(
                """SELECT s.id AS session_id, s.*, c.id AS client_id, c.prospect_id,
                          c.company_name, c.siret, c.address, c.postal_code, c.city,
                          c.contact_name, c.phone, c.email, c.folder_path,
                          t.reference AS training_reference, t.name AS training_name,
                          t.description AS training_description, t.objectives,
                          t.prerequisites, t.certification
                   FROM training_sessions s
                   JOIN clients c ON c.id=s.client_id
                   JOIN trainings t ON t.id=s.training_id
                   WHERE s.id=?""",
                (int(session_id),),
            ).fetchone()
        if row is None:
            raise ValueError("La session sélectionnée est introuvable.")
        return dict(row)

    @staticmethod
    def _format_date(value: str) -> str:
        try:
            return date.fromisoformat(value).strftime("%d/%m/%Y")
        except (TypeError, ValueError):
            return value or ""

    @staticmethod
    def _format_date_words(value: str) -> str:
        months = ("", "janvier", "février", "mars", "avril", "mai", "juin",
                  "juillet", "août", "septembre", "octobre", "novembre", "décembre")
        try:
            d = date.fromisoformat(value)
            return f"{d.day} {months[d.month]} {d.year}"
        except (TypeError, ValueError):
            return value or ""

    @classmethod
    def _format_date_range_words(cls, start_value: str, end_value: str) -> str:
        try:
            start = date.fromisoformat(start_value)
            end = date.fromisoformat(end_value)
            months = ("", "janvier", "février", "mars", "avril", "mai", "juin",
                      "juillet", "août", "septembre", "octobre", "novembre", "décembre")
            if start == end:
                return f"{start.day} {months[start.month]} {start.year}"
            if start.year == end.year and start.month == end.month:
                return f"{start.day} et {end.day} {months[end.month]} {end.year}"
            return f"du {start.day} {months[start.month]} {start.year} au {end.day} {months[end.month]} {end.year}"
        except (TypeError, ValueError):
            return f"du {start_value} au {end_value}"

    @staticmethod
    def _money(cents: int) -> str:
        return f"{int(cents) / 100:,.2f} €".replace(",", " ").replace(".00", "")

    def build_variables(self, session: dict[str, Any], document_number: str) -> dict[str, str]:
        address = " ".join(part for part in [session.get("address", ""), session.get("postal_code", ""), session.get("city", "")] if part).strip()
        start = self._format_date(str(session.get("start_date", "")))
        end = self._format_date(str(session.get("end_date", "")))
        date_range = start if start == end else f"du {start} au {end}"
        participants = int(session.get("participant_count") or 1)
        contact = str(session.get("contact_name") or "").strip()
        participant_names = contact or (f"{participants} participant" if participants == 1 else f"{participants} participants")
        participant_line = participant_names
        date_range_words = self._format_date_range_words(str(session.get("start_date", "")), str(session.get("end_date", "")))
        modality = str(session.get("modality") or "")
        location_connection = "Formation réalisée à distance via Google Meet." if "distance" in modality.lower() else modality
        variables = {
            "DOCUMENT_NUMERO": document_number,
            "DOCUMENT_DATE": datetime.now().strftime("%d/%m/%Y"),
            "CLIENT_NOM": str(session.get("company_name") or ""),
            "CLIENT_SIRET": str(session.get("siret") or ""),
            "CLIENT_ADRESSE": address,
            "CLIENT_CONTACT": str(session.get("contact_name") or ""),
            "CLIENT_TELEPHONE": str(session.get("phone") or ""),
            "CLIENT_EMAIL": str(session.get("email") or ""),
            "FORMATION_REFERENCE": str(session.get("training_reference") or ""),
            "FORMATION_NOM": str(session.get("training_name") or ""),
            "FORMATION_DESCRIPTION": str(session.get("training_description") or ""),
            "FORMATION_OBJECTIFS": str(session.get("objectives") or ""),
            "FORMATION_PREREQUIS": str(session.get("prerequisites") or ""),
            "FORMATION_CERTIFICATION": str(session.get("certification") or ""),
            "FORMATION_DUREE": f"{float(session.get('duration_hours') or 0):g} heures",
            "FORMATION_PRIX": self._money(int(session.get("price_cents") or 0)),
            "FORMATION_MODALITE": str(session.get("modality") or ""),
            "FORMATION_DATES": date_range_words,
            "FORMATION_DATES_NUMERIQUES": date_range,
            "DATE_DEBUT": start,
            "DATE_FIN": end,
            "HORAIRES": str(session.get("daily_schedule") or ""),
            "FORMATEUR_NOM": str(session.get("trainer_name") or ""),
            "COMMERCIAL_NOM": str(session.get("commercial_name") or ""),
            "FINANCEUR": str(session.get("funder") or ""),
            "FINANCEUR_DETAILS": str(session.get("funder_details") or ""),
            "PARTICIPANTS_NOMBRE": str(participants),
            "PARTICIPANTS_NOMS": participant_names,
            "PARTICIPANTS_LIGNE": participant_line,
            "OBJECTIF_PROFESSIONNEL": f"{participant_names} souhaite suivre la formation « {session.get('training_name') or ''} » afin d’acquérir et de mettre en pratique les compétences nécessaires dans le cadre de son activité professionnelle.",
            "LIEU_CONNEXION": location_connection,
            "LIEN_CONNEXION": str(session.get("connection_link") or session.get("funder_details") or "Lien communiqué avant la formation"),
            "PARTICIPANT_LIBELLE": "participant" if participants == 1 else "participants",
            "ORGANISME_NOM": "NM FORMATION",
            "ORGANISME_SIRET": "940 345 838 00012",
            "ORGANISME_ADRESSE": "347 rue Alexandra David-Néel, 59120 LOOS",
            "ORGANISME_EMAIL": "contact@forma-prof.fr",
            "ORGANISME_NDA": "32591373259",
            "ORGANISME_DIRIGEANT": "Nacim MESSADI, Président",
        }
        return variables

    def _find_template(self, document_type: str, training_id: int) -> dict[str, Any] | None:
        with self._connect() as conn:
            row = conn.execute(
                """SELECT * FROM document_templates
                   WHERE active=1 AND document_type=? AND source_path<>''
                     AND (training_id=? OR training_id IS NULL)
                   ORDER BY CASE WHEN training_id=? THEN 0 ELSE 1 END,
                            CASE WHEN name LIKE '%officiel%' THEN 0 ELSE 1 END, updated_at DESC
                   LIMIT 1""",
                (document_type, int(training_id), int(training_id)),
            ).fetchone()
        return dict(row) if row else None

    @staticmethod
    def _replace_text(text: str, variables: dict[str, str]) -> str:
        for key, value in variables.items():
            text = text.replace("{{" + key + "}}", value)
        return re.sub(r"\{\{[A-Z0-9_]+\}\}", "", text)

    def _replace_docx_variables(self, document, variables: dict[str, str]) -> None:
        def replace_paragraph(paragraph) -> None:
            original = "".join(run.text for run in paragraph.runs)
            replaced = self._replace_text(original, variables)
            if replaced != original:
                if paragraph.runs:
                    paragraph.runs[0].text = replaced
                    for run in paragraph.runs[1:]:
                        run.text = ""
                else:
                    paragraph.add_run(replaced)

        for paragraph in document.paragraphs:
            replace_paragraph(paragraph)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_paragraph(paragraph)
        for section in document.sections:
            for container in (section.header, section.footer):
                for paragraph in container.paragraphs:
                    replace_paragraph(paragraph)
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_paragraph(paragraph)

    def _merge_docx_template(self, source: Path, target: Path, variables: dict[str, str]) -> None:
        """Fusionne les balises directement dans l'OOXML sans reconstruire le DOCX.

        Les tableaux, images, styles, marges, sauts de page, en-têtes, pieds de page
        et zones de texte du modèle restent inchangés.
        """
        target.parent.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(source, "r") as zin, tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            temp_path = Path(tmp.name)
        try:
            with zipfile.ZipFile(source, "r") as zin, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                        try:
                            root = etree.fromstring(data)
                            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                            changed = False
                            ns["a"] = "http://schemas.openxmlformats.org/drawingml/2006/main"
                            paragraphs = root.xpath(".//a:p | .//w:p[not(.//a:p)]", namespaces=ns)
                            for paragraph in paragraphs:
                                nodes = paragraph.xpath(".//w:t | .//a:t", namespaces=ns)
                                if not nodes:
                                    continue
                                original = "".join(node.text or "" for node in nodes)
                                replaced = self._replace_text(original, variables)
                                if replaced != original:
                                    nodes[0].text = replaced
                                    for node in nodes[1:]:
                                        node.text = ""
                                    changed = True
                            if changed:
                                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
                        except Exception:
                            pass
                    zout.writestr(item, data)
            shutil.move(str(temp_path), str(target))
        finally:
            temp_path.unlink(missing_ok=True)

    def _base_document(self, title: str, number: str):
        if Document is None:
            raise RuntimeError("La dépendance python-docx est absente. Exécutez : pip install python-docx")
        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Inches(0.55)
        sec.bottom_margin = Inches(0.55)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Form@Prof")
        run.bold = True; run.font.size = Pt(24); run.font.color.rgb = RGBColor(51, 140, 228)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.bold = True; run.font.size = Pt(18); run.font.color.rgb = RGBColor(20, 35, 60)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(number).bold = True
        doc.add_paragraph()
        return doc

    @staticmethod
    def _add_heading(doc, text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True; run.font.size = Pt(12); run.font.color.rgb = RGBColor(51, 140, 228)

    @staticmethod
    def _add_key_values(doc, rows: Iterable[tuple[str, str]]) -> None:
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value or "—"
            cells[0].paragraphs[0].runs[0].bold = True

    def _create_default_document(self, document_type: str, variables: dict[str, str]):
        doc = self._base_document(document_type, variables["DOCUMENT_NUMERO"])
        self._add_heading(doc, "Organisme de formation")
        self._add_key_values(doc, [("Organisme", variables["ORGANISME_NOM"]), ("SIRET", variables["ORGANISME_SIRET"]),
                                   ("Adresse", variables["ORGANISME_ADRESSE"]), ("NDA", variables["ORGANISME_NDA"])])
        self._add_heading(doc, "Client")
        self._add_key_values(doc, [("Entreprise", variables["CLIENT_NOM"]), ("SIRET", variables["CLIENT_SIRET"]),
                                   ("Adresse", variables["CLIENT_ADRESSE"]), ("Contact", variables["CLIENT_CONTACT"]),
                                   ("E-mail", variables["CLIENT_EMAIL"]), ("Téléphone", variables["CLIENT_TELEPHONE"])])
        self._add_heading(doc, "Action de formation")
        self._add_key_values(doc, [("Intitulé", variables["FORMATION_NOM"]), ("Référence", variables["FORMATION_REFERENCE"]),
                                   ("Dates", variables["FORMATION_DATES"]), ("Horaires", variables["HORAIRES"]),
                                   ("Durée", variables["FORMATION_DUREE"]), ("Modalité", variables["FORMATION_MODALITE"]),
                                   ("Formateur", variables["FORMATEUR_NOM"]), ("Participants", variables["PARTICIPANTS_NOMBRE"]),
                                   ("Financeur", variables["FINANCEUR"]), ("Prix net de TVA", variables["FORMATION_PRIX"])])
        if document_type == "Devis":
            self._add_heading(doc, "Proposition financière")
            doc.add_paragraph(f"Forfait de formation : {variables['FORMATION_PRIX']} net de TVA.")
            doc.add_paragraph("TVA exonérée conformément à l’article 261-4-4° du CGI.")
            doc.add_paragraph("Devis valable 30 jours. Règlement par virement bancaire selon les conditions convenues.")
            doc.add_paragraph("Bon pour accord — Date, cachet et signature du client :")
        elif document_type == "Convention":
            self._add_heading(doc, "Objet de la convention")
            doc.add_paragraph(f"NM FORMATION organise l’action « {variables['FORMATION_NOM']} » pour {variables['CLIENT_NOM']}.")
            self._add_heading(doc, "Objectifs pédagogiques")
            doc.add_paragraph(variables["FORMATION_OBJECTIFS"] or "Les objectifs pédagogiques sont définis dans le programme annexé.")
            self._add_heading(doc, "Dispositions financières")
            doc.add_paragraph(f"Le coût total de l’action est fixé à {variables['FORMATION_PRIX']} net de TVA.")
            self._add_heading(doc, "Suivi et sanction")
            doc.add_paragraph("Le suivi est assuré par feuilles d’émargement, évaluations et attestation de fin de formation.")
            doc.add_paragraph("Fait en double exemplaire. Signatures de l’entreprise cliente et de NM FORMATION.")
        elif document_type == "Programme":
            self._add_heading(doc, "Objectifs")
            doc.add_paragraph(variables["FORMATION_OBJECTIFS"] or "Objectifs à compléter dans le catalogue de formations.")
            self._add_heading(doc, "Public et prérequis")
            doc.add_paragraph(variables["FORMATION_PREREQUIS"] or "Prérequis à compléter dans le catalogue de formations.")
            self._add_heading(doc, "Contenu")
            doc.add_paragraph(variables["FORMATION_DESCRIPTION"] or "Programme détaillé à compléter dans le catalogue de formations.")
            self._add_heading(doc, "Méthodes, suivi et évaluation")
            doc.add_paragraph("Apports théoriques, démonstrations, exercices pratiques, mise en situation et évaluation finale.")
        elif document_type == "Convocation":
            self._add_heading(doc, "Convocation")
            doc.add_paragraph(f"Madame, Monsieur,\n\nNous vous confirmons votre inscription à la formation « {variables['FORMATION_NOM']} ».")
            doc.add_paragraph(f"La formation se déroulera {variables['FORMATION_DATES']}, selon les horaires {variables['HORAIRES']}, en modalité {variables['FORMATION_MODALITE']}.")
            doc.add_paragraph("Merci de vous munir d’un ordinateur et des accès nécessaires. Les informations de connexion seront communiquées avant la session.")
        footer = doc.sections[0].footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "NM FORMATION — contact@forma-prof.fr — www.forma-prof.fr"
        return doc

    def _convert_pdf(self, docx_path: Path) -> Path | None:
        pdf_path = docx_path.with_suffix(".pdf")
        libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
        if libreoffice:
            proc = subprocess.run([libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
                                  capture_output=True, text=True, timeout=90)
            if proc.returncode == 0 and pdf_path.exists():
                return pdf_path
        if shutil.which("powershell"):
            script = (
                "$w=New-Object -ComObject Word.Application; $w.Visible=$false; "
                f"$d=$w.Documents.Open('{str(docx_path).replace("'", "''")}'); "
                f"$d.SaveAs([ref]'{str(pdf_path).replace("'", "''")}',[ref]17); $d.Close(); $w.Quit();"
            )
            try:
                proc = subprocess.run(["powershell", "-NoProfile", "-Command", script], capture_output=True, text=True, timeout=90)
                if proc.returncode == 0 and pdf_path.exists():
                    return pdf_path
            except (OSError, subprocess.TimeoutExpired):
                pass
        return None

    def generate(self, *, session_id: int, document_types: Iterable[str], generate_pdf: bool = True) -> list[GeneratedDocument]:
        if Document is None:
            raise RuntimeError("python-docx est requis pour générer les documents.")
        requested = list(dict.fromkeys(document_types))
        preflight = self.preflight(session_id, requested)
        if not preflight.can_generate:
            raise ValueError("Génération impossible :\n- " + "\n- ".join(preflight.errors))
        session = self.get_session_context(session_id)
        client_folder = Path(session.get("folder_path") or (self.project_folder / "Clients" / session["company_name"]))
        results: list[GeneratedDocument] = []
        now = datetime.now().isoformat(timespec="seconds")
        for document_type in requested:
            number = self.document_service.next_number(document_type)
            variables = self.build_variables(session, number)
            target_dir = client_folder / self.FOLDERS[document_type]
            target_dir.mkdir(parents=True, exist_ok=True)
            safe_company = re.sub(r"[^A-Za-z0-9_-]+", "_", str(session["company_name"])).strip("_") or "CLIENT"
            docx_path = target_dir / f"{number}_{safe_company}.docx"
            template = self._find_template(document_type, int(session["training_id"]))
            template_id = None
            template_name = "Modèle interne Form@Prof"
            warning = ""
            source = Path(str(template.get("source_path") or "")) if template else None
            if template and source and source.suffix.lower() == ".docx" and source.is_file():
                try:
                    self._merge_docx_template(source, docx_path, variables)
                    template_id = int(template["id"])
                    template_name = str(template.get("name") or source.name)
                except Exception as exc:
                    raise RuntimeError(f"Impossible de fusionner le modèle {document_type} : {exc}") from exc
            elif document_type == "Devis":
                doc = self._create_default_document(document_type, variables)
                doc.save(docx_path)
                template_name = "Modèle interne Form@Prof"
                warning = "Modèle interne utilisé uniquement pour compatibilité. Utilisez le workflow Demande de devis pour la production réelle."
            else:
                raise ValueError(f"Aucun modèle Word officiel valide n’est configuré pour : {document_type}.")
            if not docx_path.is_file() or docx_path.stat().st_size == 0:
                raise RuntimeError(f"Le fichier {document_type} n’a pas pu être créé : {docx_path}")
            pdf_path = self._convert_pdf(docx_path) if generate_pdf else None
            if generate_pdf and pdf_path is None:
                warning = (warning + " " if warning else "") + "PDF non créé : Word ou LibreOffice indisponible, ou conversion échouée."
            with self._connect() as conn:
                cur = conn.execute(
                    """INSERT INTO documents(document_number, document_type, prospect_id, training_id,
                       template_id, file_path, status, created_at, updated_at)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (number, document_type, session.get("prospect_id"), session.get("training_id"), template_id,
                     str(pdf_path or docx_path), "Généré", now, now),
                )
                document_id = int(cur.lastrowid)
            results.append(GeneratedDocument(document_id, document_type, number, docx_path, pdf_path, template_name, warning))
        with self._connect() as conn:
            conn.execute("UPDATE client_workflows SET current_step='Documents générés', progress_percent=35, updated_at=? WHERE session_id=?",
                         (now, int(session_id)))
        return results
